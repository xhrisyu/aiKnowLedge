"""
@Author: Chris
@Time: 2024/07/09 10:00

File converter module: convert PDF/DOCX files to images/text/markdown

1. Convert PDF to images
2. Convert PDF to text
3. Convert PDF to images and then to text
4. Convert DOCX to Markdown
5. Convert DOCX to text
Other. Exam paper processing
"""

from typing import Optional

from tqdm import tqdm
import io
import os
import json
import re
import base64
from pathlib import Path
import subprocess
import cv2
import numpy as np
from pdf2image import convert_from_path
from PIL import Image, ImageEnhance, ImageFilter
import fitz
import pytesseract
from bs4 import BeautifulSoup, NavigableString, Comment

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

from .tools import get_file_name

SUPPORTED_IMAGE_TYPES = ['png', 'jpeg']


def docx2markdown(
        docx_file_abs_path: str,
        output_dir: str
):
    def get_markdown_from_command(command):
        command_result = subprocess.run(command, shell=True, capture_output=True, text=True)
        if command_result.returncode != 0:
            raise Exception(f"Command failed: {command_result.stderr}")
        return command_result.stdout

    def convert_to_markdown_header(input_string: str) -> str:
        def replace_with_header(_match):
            header = _match.group(0).strip()
            header_no = _match.group(1).strip()
            header_text = _match.group(3).strip()

            # Check if the header contains any of the following punctuation
            check_items = [",", "，", "。", ";", "；"]
            if any(item in header for item in check_items):
                return header
            if ":" in header[:-1] or "：" in header[:-1]:
                return header

            # Check if the header is too long
            if len(header) >= 20:
                return header

            # Get the header level
            numbers = re.findall(r'\d+', header_no)
            if len(numbers) <= 3:  # Header 1/2/3
                header_level = len(numbers) * '#'
            else:  # No change for other headers
                return header

            return f"\n{header_level} {header_text}\n"

        title_pattern = r'^\s*\**\s*(\d+(\.\d+)*\.?)\s*\**\s*(.+?)\s*\**\s*$'

        return re.sub(title_pattern, replace_with_header, input_string, flags=re.MULTILINE)

    def convert_html_table_to_markdown_pipe_table(html_str: str) -> str:
        # Parse the HTML using BeautifulSoup
        _soup = BeautifulSoup(html_str, 'html.parser')

        # Remove the outermost table
        first_table = _soup.find('table')
        if first_table:
            first_table.unwrap()
        html_str = str(_soup)

        _soup = BeautifulSoup(html_str, 'html.parser')

        # Find the all table
        tables = _soup.find_all('table')

        # for table in tables:
        #     markdown_table = []
        # 
        #     # Process table headers
        #     headers = table.find('tr').find_all('th')
        #     header_row = '| ' + ' | '.join(header.get_text(strip=True) for header in headers) + ' |'
        #     markdown_table.append(header_row)
        #     markdown_table.append('|' + ' --- |' * len(headers))
        # 
        #     # Process table rows
        #     for _row in table.find_all('tr')[1:]:  # skip the header row
        #         cols = _row.find_all(['td', 'th'])  # can be 'td' or 'th' if the data rows use 'th'
        #         row_text = '| ' + ' | '.join(col.get_text(strip=True) for col in cols) + ' |'
        #         markdown_table.append(row_text)
        # 
        #     markdown_table_string = "\n".join(markdown_table)
        #     markdown_table_string_nav = NavigableString(markdown_table_string)
        #     table.replace_with("\n\n" + markdown_table_string_nav + "\n\n")
        for table in tables:
            markdown_table = []

            # Process table headers
            headers = table.find('tr').find_all('th')
            header_row = '|' + '|'.join(header.get_text(strip=True) for header in headers) + '|'
            markdown_table.append(header_row)
            markdown_table.append('|' + '---|' * len(headers))

            # Process table rows
            for _row in table.find_all('tr')[1:]:  # skip the header row
                cols = _row.find_all(['td', 'th'])  # can be 'td' or 'th' if the data rows use 'th'
                row_text = '|' + '|'.join(col.get_text(strip=True) for col in cols) + '|'
                markdown_table.append(row_text)

            markdown_table_string = "\n".join(markdown_table)
            markdown_table_string_nav = NavigableString(markdown_table_string)
            table.replace_with("\n\n" + markdown_table_string_nav + "\n\n")

        return str(_soup)

    def remove_empty_tags(html_str) -> str:
        _soup = BeautifulSoup(html_str, 'html.parser')

        # Iterate over all tags
        for tag in _soup.find_all():
            if (tag.name not in ['br', 'hr', 'img'] and  # These tags are self-closing and should not be removed
                    not tag.contents or  # Checks if the tag is empty
                    (len(tag.contents) == 1 and  # Checks if the tag contains only one item
                     isinstance(tag.contents[0], Comment))):  # Ensures it's not just a comment inside
                tag.decompose()  # Remove the tag from the soup

        return str(_soup)

    def remove_tags_with_new_line(html_str, tags_to_unwrap: list[str]) -> str:
        replaced_str = html_str
        for tag in tags_to_unwrap:
            replaced_str = replaced_str.replace(f"<{tag}>", "\n\n")
            replaced_str = replaced_str.replace(f"</{tag}>", "\n\n")
        return replaced_str

    def replace_img_tags(html_str: str) -> str:

        _soup = BeautifulSoup(html_str, 'html.parser')

        # Find all <img> tags and replace them with Markdown format
        for img in _soup.find_all('img'):
            src = img.get('src')
            alt = img.get('alt', '')
            # Create the markdown replacement string
            markdown_img = f"![{alt}]({src})"
            # Replace the <img> tag with a NavigableString containing the markdown
            img.replace_with(NavigableString(markdown_img))

        return str(_soup)

    def iteratively_unwrap_tags_with_regex(html_str: str, tags_to_unwrap=None) -> str:
        if tags_to_unwrap is None:
            tags_to_unwrap = ['p', 'tbody', 'tr', 'td']
        regex_pattern = fr'<({"|".join(tags_to_unwrap)})\b[^>]*>(.*?)</\1>'
        previous_html = ""

        # Continue looping until no changes are made to the HTML string
        while previous_html != html_str:
            previous_html = html_str
            html_str = re.sub(regex_pattern, r'\2', html_str, flags=re.DOTALL)

        return html_str

    def clean_text(input_string: str) -> str:
        input_string = re.sub(r'\n{3,}', '\n\n', input_string)  # 替换多个连续换行(>=3)为2个换行
        input_string = re.sub(r' +', ' ', input_string)  # 替换多个连续空格为1个空格
        return input_string

    def extract_and_save_images(markdown_text: str, root_output_dir: str, file_name: str):

        # Match base64 image string
        # Define image output folder
        image_output_dir = os.path.join(root_output_dir, "image")
        Path(root_output_dir).mkdir(parents=True, exist_ok=True)
        Path(image_output_dir).mkdir(parents=True, exist_ok=True)

        image_pattern = re.compile(r'(["(])data:image/(png|jpeg|x-emf);base64,([^")]+)([")])')
        image_matches = image_pattern.findall(markdown_text)

        # Process each image matches
        for i, (quote_start, image_type, base64_data, quote_end) in enumerate(image_matches):
            image_data = base64.b64decode(base64_data)  # Decode base64 image

            image_file_path = f"{image_output_dir}/image_{i}.{image_type}"
            image_file_relative_path = f"./image/image_{i}.{image_type}"

            with open(image_file_path, "wb") as file:
                file.write(image_data)

            markdown_text = markdown_text.replace(
                f'{quote_start}data:image/{image_type};base64,{base64_data}{quote_end}',
                f'{quote_start}{image_file_relative_path}{quote_end}'
            )

        # Process title level
        markdown_text = convert_to_markdown_header(markdown_text)

        # Replace img tag with Markdown format
        markdown_text = replace_img_tags(markdown_text)

        # Convert html tags to Markdown format
        markdown_text = convert_html_table_to_markdown_pipe_table(markdown_text)

        # Regular expression to remove tags but keep the text inside them
        markdown_text = iteratively_unwrap_tags_with_regex(
            markdown_text,
            tags_to_unwrap=['p', 'strong', 'tbody', 'tr', 'td', 'th', 'li']
        )

        # Remove <a></a> tags and their contents
        markdown_text = re.sub(r'<a[^>]*>(.*?)</a>', '', markdown_text)

        # Regular expression to remove 'alt="..."' attributes from <img> tags
        markdown_text = re.sub(r'(<img[^>]*?)\s*alt="[^"]*"([^>]*>)', r'\1\2', markdown_text)

        # Remove <ol></ol> <ul></ul> with text and add new line
        markdown_text = remove_tags_with_new_line(markdown_text, tags_to_unwrap=['ol', 'ul'])

        # Remove void tag
        markdown_text = remove_empty_tags(markdown_text)

        # Replace multiple spaces/newlines
        markdown_text = clean_text(markdown_text)

        # Save markdown file
        with open(os.path.join(root_output_dir, f"{file_name}.md"), "w", encoding="utf-8") as file:
            file.write(markdown_text)

    # Convert DOCX to Markdown
    w2m_command = f'source ~/.bash_profile && w2m "{docx_file_abs_path}"'
    markdown_code = get_markdown_from_command(w2m_command)

    extract_and_save_images(
        markdown_text=markdown_code,
        root_output_dir=output_dir,
        file_name=get_file_name(docx_file_abs_path, with_extension=False)
    )


def pdf2image(
        pdf_path: str,
        output_path: str = None,
        image_type: str = 'png',
        return_images: bool = True
) -> list[Image.Image] | list[str]:
    """
    Convert a PDF file to images.

    :param pdf_path: original PDF file path
    :param output_path: output image path (if return_images is False, the images will be saved to this path)
    :param image_type: image type, 'png' or 'jpeg'
    :param return_images: if True, return the images, else save the images to the output path
    :return: list of images / list of image paths
    """

    def enhance_and_dilate(_image):
        # Convert to grayscale image
        gray_image = cv2.cvtColor(np.array(_image), cv2.COLOR_BGR2GRAY)

        # Binarize the image
        _, binary_image = cv2.threshold(gray_image, 128, 255, cv2.THRESH_BINARY_INV)

        # Apply dilation to make black areas thicker
        kernel = np.ones((2, 2), np.uint8)
        dilated_image = cv2.dilate(binary_image, kernel, iterations=1)

        # Invert the binary image
        dilated_image = cv2.bitwise_not(dilated_image)

        # Convert back to PIL image
        return Image.fromarray(dilated_image)

    # Check image_type
    if image_type not in SUPPORTED_IMAGE_TYPES:
        raise ValueError(f"Unsupported image type: {image_type}, supported types are 'png' and 'jpeg'")

    # Convert pdf pages to images
    raw_images = convert_from_path(pdf_path, dpi=200, fmt=image_type)

    # Process images
    images = []
    for image in raw_images:
        # Raise up contrast(对比度)
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(3)

        # Raise up brightness(亮度)
        enhancer = ImageEnhance.Brightness(image)
        image = enhancer.enhance(1.5)

        # Sharpen the image(锐化)
        image = image.filter(ImageFilter.SHARPEN)

        # Enhance and dilate the image(增强并膨胀)
        image = enhance_and_dilate(image)
        images.append(image)

    # if `return_images` is True, return the images
    if return_images:
        return images

    # else save the images to the output path, and return the image paths
    # Check the validity of `output_path`
    if output_path is None:
        output_path = os.path.join(os.getcwd(), 'output_pdf_images')
    if not os.path.exists(output_path):
        os.makedirs(output_path)

    image_paths = []
    base_file_name = get_file_name(pdf_path, with_extension=False)
    for page_idx, image in enumerate(images):
        cur_image_path = os.path.join(output_path, f'{base_file_name}_p{page_idx + 1}.{image_type}')
        cur_image_path = os.path.abspath(cur_image_path)
        image.save(cur_image_path, format=image_type)
        image_paths.append(cur_image_path)

    return image_paths


def pdf2txt(pdf_input_path: str, txt_output_path: Optional[str]) -> str:
    # Using fitz
    with fitz.open(pdf_input_path) as pdf:
        text = ""
        for page in tqdm(pdf):
            text += page.get_text()

    # Extract the file content
    lines = text.split('\n')

    # Regular expression for matching Chinese characters and English words
    chinese_chars_pattern = r'[\u4e00-\u9fff]+'
    english_words_pattern = r'[A-Za-z]+'

    # Process each line
    processed_lines = []
    for line in lines:
        line = line.strip()
        line = re.sub(r'\s+', ' ', line)

        # Remove spaces between Chinese characters
        line = re.sub(chinese_chars_pattern, lambda x: x.group().replace(' ', ''), line)
        # Ensure only one space between English words
        line = re.sub(english_words_pattern, lambda x: ' '.join(x.group().split()), line)

        if line != '':
            processed_lines.append(line)

    # Write the processed lines to the output file
    if txt_output_path:
        with open(txt_output_path, 'w', encoding='utf-8') as file:
            for line in processed_lines:
                file.write(line + '\n')

    return "\n".join(processed_lines)


def pdf2img2txt(pdf_input_path: str, txt_output_path: str):
    doc = fitz.open(pdf_input_path)
    text = ""
    for page_num in range(len(doc)):
        # load the page
        page = doc.load_page(page_num)

        # get the page as a pixmap and convert to bytes
        pix = page.get_pixmap()
        img_data = pix.tobytes("png")

        # using opencv to convert image to text
        image = Image.open(io.BytesIO(img_data))
        cur_text = pytesseract.image_to_string(image, lang='chi_sim')
        print(cur_text)
        text += cur_text

    doc.close()

    with open(txt_output_path, 'w', encoding='utf-8') as file:
        file.write(text)


def docx2txt(docx_input_path: str, txt_output_path: str) -> None:
    def iter_block_items(parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # Load the docx file
    doc = Document(f"{docx_input_path}")

    # Extract the text content
    text_content = []
    for block in iter_block_items(doc):
        # 1. read Paragraph
        if isinstance(block, Paragraph):
            para_text = block.text.replace(" ", "")
            if para_text != "":
                # print(para_text)
                text_content.append(para_text)

        # 2. read table
        elif isinstance(block, Table):
            skip_flag = False
            for row in block.rows:
                pre_cell_text = ""
                if skip_flag:
                    break
                for cell in row.cells:
                    cell_text = (
                        "."
                        .join([paragraph.text for paragraph in cell.paragraphs])
                        .replace("\n", ".").replace(" ", "")
                    )
                    if cell_text == "更改标记":
                        skip_flag = True
                        break
                    if cell_text == pre_cell_text or cell_text == "":  # 去重
                        continue
                    pre_cell_text = cell_text
                    text_content.append(cell_text)
                # print(pre_cell_text)

    # Save the text content to the txt file
    with open(txt_output_path, "w", encoding="utf-8") as txt_file:
        for text in text_content:
            text = text.replace(" ", "").replace("\u00A0", "").replace("\n", "")
            if text != "":
                txt_file.write(text + "\n")


class ExamPaperProcessor:
    @staticmethod
    def process_raw_text_to_dict(raw_text):
        # 使用正则表达式提取题目信息
        match = re.match(r'(\d+)、(.*?)<(\w+)>(.*?)\n', raw_text)
        if match:
            question_id = int(match.group(1))
            question_text = match.group(2) + "<>" + match.group(4)
            answers = match.group(3)

            # 使用正则表达式提取选项信息
            # choices_match = re.findall(r'([A-D])、(.*?)\n', input_string)
            choices_match = re.findall(r'([A-D])、(.*?)(?=[A-D]、|\n|$)', raw_text)
            choices = {choice[0]: choice[1] for choice in choices_match}

            # 构建字典
            result = {
                'id': question_id,
                # 'question': question_text.replace('<>', '<{}>'.format(answer)),
                'question': question_text,
                'choices': choices,
                'answer': [answer for answer in answers]
            }
            return result
        else:
            return None

    @staticmethod
    def preprocess_exam_paper(directory_path: str = './document_exam'):
        file_names = [file_name for file_name in os.listdir(directory_path) if file_name.endswith(".docx")]
        output_directory_path = './document_exam_json'
        for file_name in file_names:
            print(f"processing {file_name}...")
            doc = Document(f"{directory_path}/{file_name}")
            output_file_name = file_name.split(".")[0]

            text = "".join([paragraph.text + "\n" for paragraph in doc.paragraphs]).replace(" ", "")
            # 替换括号
            text = text.replace("(", "<").replace(")", ">").replace("（", "<").replace("）", ">")

            # 使用正则表达式分割文本成题目段落
            pattern = r'\d+、.*?(?=\d+、|\Z)'
            questions = re.findall(pattern, text, re.DOTALL)

            # 打印分段后的题目段落
            json_data = []
            for index, question in enumerate(questions, start=1):
                print(f'题目 {index}:')
                json_question = ExamPaperProcessor.process_raw_text_to_dict(question)
                json_data.append(json_question)
                print(json_question)
                print('-' * 50)

            with open(f"{output_directory_path}/{output_file_name}.json", "w") as json_file:
                json.dump(json_data, json_file, ensure_ascii=False, indent=4)
