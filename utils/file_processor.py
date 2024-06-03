import io
import re
import os
# from pdfminer.high_level import extract_text
import pytesseract
import fitz
from PIL import Image
from tqdm import tqdm
import json

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

from langchain_community.document_loaders import PyPDFium2Loader
from utils.tools import get_file_path_list


class PDFProcessor:
    @staticmethod
    def pdf_to_txt(pdf_input_path: str, txt_output_path: str | None = None) -> str:
        # Using fitz
        with fitz.open(pdf_input_path) as pdf:
            text = ""
            for page in tqdm(pdf):
                text += page.get_text()

        # Extract the file content
        lines = text.split('\n')

        # Regular expression for matching Chinese characters and English words
        chinese_chars_pattern = r'[\u4e00-\u9fff]+'
        english_words_pattern = r'[A-Za-z]+'

        # Process each line
        processed_lines = []
        for line in lines:
            line = line.strip()
            line = re.sub(r'\s+', ' ', line)

            # Remove spaces between Chinese characters
            line = re.sub(chinese_chars_pattern, lambda x: x.group().replace(' ', ''), line)
            # Ensure only one space between English words
            line = re.sub(english_words_pattern, lambda x: ' '.join(x.group().split()), line)

            if line != '':
                processed_lines.append(line)

        # Write the processed lines to the output file
        if txt_output_path:
            with open(txt_output_path, 'w', encoding='utf-8') as file:
                for line in processed_lines:
                    file.write(line + '\n')

        return "\n".join(processed_lines)

    @staticmethod
    def pdf_to_img_to_txt(pdf_input_path: str, txt_output_path: str):
        doc = fitz.open(pdf_input_path)
        text = ""
        for page_num in range(len(doc)):
            # load the page
            page = doc.load_page(page_num)

            # get the page as a pixmap and convert to bytes
            pix = page.get_pixmap()
            img_data = pix.tobytes("png")

            # using opencv to convert image to text
            image = Image.open(io.BytesIO(img_data))
            cur_text = pytesseract.image_to_string(image, lang='chi_sim')
            print(cur_text)
            text += cur_text

        doc.close()

        with open(txt_output_path, 'w', encoding='utf-8') as file:
            file.write(text)


class DocxProcessor:
    @staticmethod
    def docx_to_txt(docx_input_path: str, txt_output_path: str) -> None:
        def iter_block_items(parent):
            if isinstance(parent, _Document):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            elif isinstance(parent, _Row):
                parent_elm = parent._tr
            else:
                raise ValueError("something's not right")
            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

        # Load the docx file
        doc = Document(f"{docx_input_path}")

        # Extract the text content
        text_content = []
        for block in iter_block_items(doc):
            # 1. read Paragraph
            if isinstance(block, Paragraph):
                para_text = block.text.replace(" ", "")
                if para_text != "":
                    # print(para_text)
                    text_content.append(para_text)

            # 2. read table
            elif isinstance(block, Table):
                skip_flag = False
                for row in block.rows:
                    pre_cell_text = ""
                    if skip_flag:
                        break
                    for cell in row.cells:
                        cell_text = (
                            "."
                            .join([paragraph.text for paragraph in cell.paragraphs])
                            .replace("\n", ".").replace(" ", "")
                        )
                        if cell_text == "更改标记":
                            skip_flag = True
                            break
                        if cell_text == pre_cell_text or cell_text == "":  # 去重
                            continue
                        pre_cell_text = cell_text
                        text_content.append(cell_text)
                    # print(pre_cell_text)

        # Save the text content to the txt file
        with open(txt_output_path, "w", encoding="utf-8") as txt_file:
            for text in text_content:
                text = text.replace(" ", "").replace("\u00A0", "").replace("\n", "")
                if text != "":
                    txt_file.write(text + "\n")


class ExamPaperProcessor:
    @staticmethod
    def process_raw_text_to_dict(raw_text):
        # 使用正则表达式提取题目信息
        match = re.match(r'(\d+)、(.*?)<(\w+)>(.*?)\n', raw_text)
        if match:
            question_id = int(match.group(1))
            question_text = match.group(2) + "<>" + match.group(4)
            answers = match.group(3)

            # 使用正则表达式提取选项信息
            # choices_match = re.findall(r'([A-D])、(.*?)\n', input_string)
            choices_match = re.findall(r'([A-D])、(.*?)(?=[A-D]、|\n|$)', raw_text)
            choices = {choice[0]: choice[1] for choice in choices_match}

            # 构建字典
            result = {
                'id': question_id,
                # 'question': question_text.replace('<>', '<{}>'.format(answer)),
                'question': question_text,
                'choices': choices,
                'answer': [answer for answer in answers]
            }
            return result
        else:
            return None

    @staticmethod
    def preprocess_exam_paper(directory_path: str = './document_exam'):
        file_names = [file_name for file_name in os.listdir(directory_path) if file_name.endswith(".docx")]
        output_directory_path = './document_exam_json'
        for file_name in file_names:
            print(f"processing {file_name}...")
            doc = Document(f"{directory_path}/{file_name}")
            output_file_name = file_name.split(".")[0]

            text = "".join([paragraph.text + "\n" for paragraph in doc.paragraphs]).replace(" ", "")
            # 替换括号
            text = text.replace("(", "<").replace(")", ">").replace("（", "<").replace("）", ">")

            # 使用正则表达式分割文本成题目段落
            pattern = r'\d+、.*?(?=\d+、|\Z)'
            questions = re.findall(pattern, text, re.DOTALL)

            # 打印分段后的题目段落
            json_data = []
            for index, question in enumerate(questions, start=1):
                print(f'题目 {index}:')
                json_question = ExamPaperProcessor.process_raw_text_to_dict(question)
                json_data.append(json_question)
                print(json_question)
                print('-' * 50)

            with open(f"{output_directory_path}/{output_file_name}.json", "w") as json_file:
                json.dump(json_data, json_file, ensure_ascii=False, indent=4)


def remove_redundant_repetitions(s):
    s = re.sub(r'\s+', ' ', s)  # 替换连续的空格为单个空格
    # 尝试找出重复的短字符串，假设短字符串的长度至少为 n
    for n in range(len(s) // 2, 0, -1):
        # 生成长度为 n 的短字符串的正则表达式
        pattern = r'(.{' + str(n) + r'})(?:\s+\1)+'
        match = re.search(pattern, s)
        if match:
            # 如果找到重复的短字符串，则返回其中一个
            return match.group(1).strip()

    return s.strip()


def process_pdf(folder_path: str, output_folder_path: str) -> None:

    # Define the output folder path
    if not os.path.exists(f"{output_folder_path}"):
        os.makedirs(f"{output_folder_path}")

    # Get absolute path
    folder_path = os.path.abspath(folder_path)
    output_folder_path = os.path.abspath(output_folder_path)

    # Get all the PDF file names
    pdf_file_path_list = get_file_path_list(root_path=folder_path, file_suffix="pdf")

    # Process each PDF file
    for pdf_file_path in tqdm(pdf_file_path_list):
        print(f"processing {pdf_file_path}...")

        # Get the file name without suffix
        pdf_file_name = os.path.basename(pdf_file_path)
        pdf_file_name_without_suffix, _ = os.path.splitext(pdf_file_name)

        # Process the PDF file
        PDFProcessor.pdf_to_txt(
            pdf_input_path=pdf_file_path,
            txt_output_path=f"{output_folder_path}/{pdf_file_name_without_suffix}.txt"
        )
        # PDFFileProcessor.pdf_to_img_to_txt(
        #     pdf_input_path=pdf_file_path,
        #     txt_output_path=f"{output_folder_path}/{pdf_file_name_without_suffix}.txt"
        # )


def process_pdf_langchain(folder_path: str, output_folder_path: str) -> None:
    # PyPDFium2

    # Define the output folder path
    if not os.path.exists(f"{output_folder_path}"):
        os.makedirs(f"{output_folder_path}")

    # Get absolute path
    folder_path = os.path.abspath(folder_path)
    output_folder_path = os.path.abspath(output_folder_path)

    # Get all the PDF file names
    pdf_file_path_list = get_file_path_list(root_path=folder_path, file_suffix="pdf")

    # Process each PDF file
    for pdf_file_path in tqdm(pdf_file_path_list):
        print(f"processing {pdf_file_path}...")

        # Load PDF file
        pdf_loader = PyPDFium2Loader(pdf_file_path)
        pdf_data = pdf_loader.load()

        # Process each pdf file page
        pdf_length = len(pdf_data)
        pdf_content = ""
        for page_num in range(pdf_length):
            page_text = pdf_data[page_num].page_content
            sentences = []
            buffer = ''
            page_text_list = page_text.split('\r\n')
            for index in range(len(page_text_list)):
                cur_text = page_text_list[index]
                cur_text = cur_text.strip().replace('', '')
                if cur_text:
                    # The situation of duplicate title
                    cur_text = remove_redundant_repetitions(cur_text)

                    # If the length is less than 6 and is th first 2 row, it's probably a title
                    if index < 2 and len(cur_text) < 6:
                        sentences.append(cur_text)
                        continue

                    if buffer:
                        if re.search(r'[。.？?!！；;：:…】）)]$', buffer[-1]):
                            sentences.append(buffer)
                            if (re.search(r'^[\x01●□√（(【[“‘《·0-9]', cur_text[:4]) and len(cur_text) <= 20) or re.search(
                                    r'[。.？?!！；;：:…】）)]$', cur_text[-1]):
                                sentences.append(cur_text)
                                buffer = ""
                            else:
                                buffer = cur_text
                        else:
                            buffer += cur_text
                    else:
                        buffer += cur_text

            # Add the last buffer if it's not empty
            if buffer:
                sentences.append(buffer)

            # Add the processed sentences to the pdf content
            pdf_content += '\n'.join(sentences)

        # Get the file name without suffix
        pdf_file_name = os.path.basename(pdf_file_path)
        pdf_file_name_without_suffix, _ = os.path.splitext(pdf_file_name)

        # Write the processed content to the output file
        with open(f"{output_folder_path}/{pdf_file_name_without_suffix}.txt", 'w', encoding='utf-8') as file:
            file.writelines(pdf_content)


def process_docx(folder_path: str, output_folder_path: str) -> None:

    # Define the output folder path
    if not os.path.exists(f"{output_folder_path}"):
        os.makedirs(f"{output_folder_path}")

    # Get absolute path
    folder_path = os.path.abspath(folder_path)
    output_folder_path = os.path.abspath(output_folder_path)

    # Get all the docx file names
    docx_file_path = get_file_path_list(root_path=folder_path, file_suffix="docx")

    # Process each docx file
    for docx_file_path in tqdm(docx_file_path):
        print(f"processing {docx_file_path}...")

        # Get the file name without suffix
        docx_file_name = os.path.basename(docx_file_path)
        docx_file_name_without_suffix, _ = os.path.splitext(docx_file_name)

        # Process the docx file
        DocxProcessor.docx_to_txt(
            docx_input_path=docx_file_path,
            txt_output_path=f"{output_folder_path}/{docx_file_name_without_suffix}.txt"
        )


if __name__ == "__main__":
    # load_dotenv()
    # openai.api_key = os.environ.get("OPENAI_API_KEY")

    # ------ Test table content extraction ------
    # table_pic_path = os.getcwd() + "/doc_raw/图1：长度测量溯源性链的示例.jpg"
    # description = get_table_content(table_pic_path, openai.api_key)
    # print(description)

    # ------ Test image content extraction ------
    # table_pic_path = os.getcwd() + "/doc_raw/图1：单一过程要素示意图.jpg"
    # description = get_image_content(table_pic_path, openai.api_key)
    # print(description)

    # ------ Process pdf txt text ------
    # process_pdf(folder_path='doc_raw/第一期知识库part3', output_folder_path='doc/kb_4_pdf_2_txt')
    process_pdf_langchain(folder_path='../doc_raw/第一期知识库part3', output_folder_path='../doc/kb_4_pdf_2_txt')

    # ------ Process docx txt text ------
    # process_docx(folder_path='doc_raw/第一期知识库part3', output_folder_path='doc/kb_4_docx_2_txt')
